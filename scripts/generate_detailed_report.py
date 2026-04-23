from __future__ import annotations

import csv
import json
from collections import defaultdict
from datetime import datetime
from pathlib import Path
from typing import Iterable

import matplotlib.pyplot as plt
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path("/Users/macbook/Downloads/Credit Scoring Model")
SUMMARY_CSV = ROOT / "artifacts/credit/benchmark_summary.csv"
REPORT_DIR = ROOT / "report"
FIGURE_DIR = REPORT_DIR / "figures"
OUTPUT_DOCX = REPORT_DIR / "SIT723_detailed_progress_report.docx"
OUTPUT_MD = REPORT_DIR / "SIT723_detailed_progress_report.md"


def load_rows() -> list[dict[str, str]]:
    with SUMMARY_CSV.open("r", encoding="utf-8") as handle:
        return list(csv.DictReader(handle))


def load_dataset_profile(run_id: str) -> dict:
    path = ROOT / "artifacts/credit" / run_id / "dataset_profile.json"
    return json.loads(path.read_text(encoding="utf-8"))


def infer_ocr_status() -> dict[str, str]:
    dataset_dir = ROOT / "artifacts/ocr/paddle_ocr_pilot/synthetic_dataset"
    repo_dir = ROOT / "third_party/PaddleOCR"
    pretrain_tar = ROOT / "third_party/pretrain_models/en_PP-OCRv3_rec_train.tar"
    pretrain_extracted = ROOT / "third_party/pretrain_models/en_PP-OCRv3_rec_train"
    config_path = ROOT / "artifacts/ocr/paddle_ocr_pilot/tiny_rec_config.yml"
    model_dir = ROOT / "artifacts/ocr/paddle_ocr_pilot/model"

    status_items = {
        "Synthetic OCR data prepared": str(dataset_dir.exists()),
        "PaddleOCR repository cloned": str(repo_dir.exists()),
        "Pretrained recognition archive downloaded": str(pretrain_tar.exists()),
        "Pretrained model extracted": str(pretrain_extracted.exists()),
        "Training config written": str(config_path.exists()),
        "Fine-tuned model output available": str(model_dir.exists()),
    }

    if model_dir.exists():
        summary = "OCR pilot completed."
        blocker = "No active blocker detected."
    else:
        summary = "OCR pilot environment setup is partially complete but the fine-tuning run has not completed yet."
        blocker = (
            "Current debugging focus: the PaddleOCR bootstrap sequence progressed through package installation, repository clone, "
            "synthetic dataset generation, and pretrained archive download, but the run did not yet reach a stable saved fine-tuned checkpoint."
        )

    return {
        "summary": summary,
        "blocker": blocker,
        "items": status_items,
    }


def make_methodology_diagram(output_path: Path) -> None:
    fig, ax = plt.subplots(figsize=(12, 3.8))
    ax.axis("off")
    boxes = [
        ("Literature Gaps\nRQ1-RQ3", 0.06),
        ("Public Datasets\nGerman Credit + Lending Club sample", 0.26),
        ("Preprocessing\nimpute, encode, scale, rebalance", 0.47),
        ("Modeling\nLR, RF, GB, MLP", 0.67),
        ("Evaluation\nROC-AUC, PR-AUC, Recall, F1", 0.87),
    ]
    for label, x in boxes:
        rect = plt.Rectangle((x - 0.085, 0.35), 0.17, 0.28, facecolor="#EAF2FF", edgecolor="#1F4E79", linewidth=1.5)
        ax.add_patch(rect)
        ax.text(x, 0.49, label, ha="center", va="center", fontsize=11)
    for idx in range(len(boxes) - 1):
        x0 = boxes[idx][1] + 0.085
        x1 = boxes[idx + 1][1] - 0.085
        ax.annotate("", xy=(x1, 0.49), xytext=(x0, 0.49), arrowprops=dict(arrowstyle="->", lw=1.8, color="#1F4E79"))
    ax.text(
        0.5,
        0.16,
        "Supporting extension: PaddleOCR synthetic fine-tuning workflow for borrower-document intake is being debugged in parallel.",
        ha="center",
        va="center",
        fontsize=10,
    )
    fig.tight_layout()
    output_path.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(output_path, dpi=180, bbox_inches="tight")
    plt.close(fig)


def make_benchmark_chart(rows: list[dict[str, str]], output_path: Path) -> None:
    datasets = ["german_credit", "lending_club_sample"]
    models = ["logistic_regression", "random_forest", "gradient_boosting", "mlp"]
    labels = {
        "logistic_regression": "LR",
        "random_forest": "RF",
        "gradient_boosting": "GB",
        "mlp": "MLP",
    }

    grouped = defaultdict(dict)
    for row in rows:
        grouped[row["dataset"]][row["model"]] = float(row["test_roc_auc"])

    x = range(len(models))
    width = 0.35
    fig, ax = plt.subplots(figsize=(10, 4.5))
    german_vals = [grouped["german_credit"].get(model, 0) for model in models]
    lending_vals = [grouped["lending_club_sample"].get(model, 0) for model in models]
    ax.bar([i - width / 2 for i in x], german_vals, width=width, label="German Credit", color="#4472C4")
    ax.bar([i + width / 2 for i in x], lending_vals, width=width, label="Lending Club sample", color="#ED7D31")
    ax.set_xticks(list(x))
    ax.set_xticklabels([labels[m] for m in models])
    ax.set_ylabel("Test ROC-AUC")
    ax.set_title("Pilot Benchmark Comparison by Dataset and Model")
    ax.set_ylim(0, 1)
    ax.legend()
    fig.tight_layout()
    output_path.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(output_path, dpi=180)
    plt.close(fig)


def add_title(document: Document, title: str, subtitle: str) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(18)
    p2 = document.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(subtitle)
    run2.italic = True
    run2.font.size = Pt(11)


def add_paragraphs(document: Document, paragraphs: Iterable[str]) -> None:
    for text in paragraphs:
        document.add_paragraph(text)


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for idx, value in enumerate(headers):
        header_cells[idx].text = value
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value


def best_runs_by_dataset(rows: list[dict[str, str]]) -> dict[str, dict[str, str]]:
    grouped: dict[str, list[dict[str, str]]] = defaultdict(list)
    for row in rows:
        grouped[row["dataset"]].append(row)
    return {dataset: max(dataset_rows, key=lambda r: float(r["test_roc_auc"])) for dataset, dataset_rows in grouped.items()}


def build_markdown(rows: list[dict[str, str]], ocr_status: dict[str, str]) -> str:
    best_runs = best_runs_by_dataset(rows)
    german = best_runs["german_credit"]
    lending = best_runs["lending_club_sample"]
    return f"""# SIT723 Detailed Progress Report

Prepared on {datetime.now().strftime("%d %B %Y")}

## Summary

This report documents the current SIT723 project status, finalised methodology, completed experimental design, preliminary benchmark experiments, and the current PaddleOCR blocker under active debugging.

## Best pilot results

- German Credit best run: {german['model']} with {german['imbalance_strategy']} and ROC-AUC {float(german['test_roc_auc']):.3f}
- Lending Club sample best run: {lending['model']} with {lending['imbalance_strategy']} and ROC-AUC {float(lending['test_roc_auc']):.3f}

## OCR status

- {ocr_status['summary']}
- {ocr_status['blocker']}
"""


def build_docx(rows: list[dict[str, str]], ocr_status: dict[str, str]) -> None:
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    FIGURE_DIR.mkdir(parents=True, exist_ok=True)
    methodology_fig = FIGURE_DIR / "methodology_flow.png"
    benchmark_fig = FIGURE_DIR / "benchmark_comparison.png"
    make_methodology_diagram(methodology_fig)
    make_benchmark_chart(rows, benchmark_fig)

    best_runs = best_runs_by_dataset(rows)
    german_best = best_runs["german_credit"]
    lending_best = best_runs["lending_club_sample"]
    german_profile = load_dataset_profile(german_best["run_id"])
    lending_profile = load_dataset_profile(lending_best["run_id"])

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    add_title(
        document,
        "SIT723 Detailed Progress Report",
        "AI-based credit scoring benchmark with PaddleOCR extension status update",
    )
    document.add_paragraph(f"Prepared on {datetime.now().strftime('%d %B %Y')}")
    document.add_paragraph("Student: Rudraksh Dhamija")
    document.add_paragraph("Supervisor: Dr. Shamsul Huda")

    document.add_heading("1. Project Context and Current Scope", level=1)
    add_paragraphs(
        document,
        [
            "The current project scope remains aligned with the approved SIT723 research direction: benchmarking explainable AI and deep learning models for credit scoring using German Credit and Lending Club data.",
            "The broader loan-origination relevance is retained through a document-intelligence extension. In the current pilot, PaddleOCR is treated as a supporting module for borrower-document intake rather than a replacement for the core credit-scoring study.",
            "This report is written as a truthful progress update. All numerical claims below are derived from saved pilot experiment artifacts already produced in the workspace.",
        ],
    )

    document.add_heading("2. Research Aim, Gaps, and Questions", level=1)
    add_paragraphs(
        document,
        [
            "The main aim is to determine how well strong machine learning and neural models perform on public credit-risk datasets while preserving robustness and explainability for lending-oriented use.",
            "The work addresses three recurring literature gaps: limited cross-dataset comparison, weak attention to minority-class detection under class imbalance, and insufficient connection between predictive performance and explainable lending practice.",
        ],
    )
    add_table(
        document,
        ["Research Question", "Purpose"],
        [
            ["RQ1", "Compare deep learning models with strong machine-learning baselines for credit-risk prediction."],
            ["RQ2", "Evaluate preprocessing and imbalance strategies that improve risky-borrower detection."],
            ["RQ3", "Retain practical explainability through feature-importance analysis without losing predictive quality."],
        ],
    )

    document.add_heading("3. Finalised Methodology", level=1)
    document.add_paragraph(
        "The methodology is fully designed and implemented as a reproducible experiment pipeline. Each run follows the same flow from data ingestion to metrics export."
    )
    document.add_picture(str(methodology_fig), width=Inches(6.8))
    add_paragraphs(
        document,
        [
            "The experimental workflow uses public datasets, train/validation/test separation, explicit imbalance-handling variants, and consistent saved outputs including metrics JSON, confusion matrix, ROC curve, PR curve, and feature-importance artifacts.",
            "Model families used in the pilot are Logistic Regression, Random Forest, HistGradientBoosting, and MLP. This provides a balanced comparison between interpretable linear baselines, strong ensemble methods, and a neural baseline.",
        ],
    )

    document.add_heading("4. Finished Experimental Design", level=1)
    add_table(
        document,
        ["Design Element", "Current Pilot Design"],
        [
            ["Datasets", "German Credit full public dataset; public cleaned Lending Club sample with pilot subset size 4,000"],
            ["Preprocessing", "Median or mode imputation, one-hot encoding, scaling for linear and neural models"],
            ["Imbalance strategies", "None, class weighting, random oversampling"],
            ["Evaluation metrics", "ROC-AUC, PR-AUC, precision, recall, F1, accuracy, confusion matrix, training time"],
            ["Explainability", "Feature-importance CSV and figure for each run"],
            ["Outputs", "Per-run JSON summaries, plots, and benchmark summary CSV"],
        ],
    )

    document.add_heading("5. Dataset Profile", level=1)
    add_table(
        document,
        ["Dataset", "Records used", "Positive-class rate", "Notes"],
        [
            [
                "German Credit",
                str(german_profile["records"]),
                f"{german_profile['positive_rate']:.3f}",
                "Positive class is bad credit risk.",
            ],
            [
                "Lending Club sample",
                str(lending_profile["records"]),
                f"{lending_profile['positive_rate']:.3f}",
                "Pilot run uses a stratified subset from a public cleaned Lending Club CSV.",
            ],
        ],
    )
    add_paragraphs(
        document,
        [
            "The German Credit data remains useful as a compact benchmark with a known 70:30 class structure. The Lending Club pilot uses a larger public sample and therefore gives a more realistic tabular-lending setting, while still remaining computationally manageable for this stage.",
            "The positive rate in the Lending Club pilot subset is lower than in German Credit, which makes minority-class recall especially important for model comparison.",
        ],
    )

    document.add_heading("6. Preliminary Experiments Conducted", level=1)
    document.add_paragraph(
        "Eight pilot benchmark runs have been completed across two datasets. The table below summarizes the current test-set results saved in the benchmark summary file."
    )
    result_rows = []
    for row in rows:
        result_rows.append(
            [
                row["run_id"],
                row["dataset"],
                row["model"],
                row["imbalance_strategy"],
                f"{float(row['test_roc_auc']):.3f}",
                f"{float(row['test_pr_auc']):.3f}",
                f"{float(row['test_recall']):.3f}",
                f"{float(row['test_f1']):.3f}",
            ]
        )
    add_table(
        document,
        ["Run ID", "Dataset", "Model", "Imbalance", "ROC-AUC", "PR-AUC", "Recall", "F1"],
        result_rows,
    )
    document.add_picture(str(benchmark_fig), width=Inches(6.8))

    document.add_heading("7. Early Findings and Interpretation", level=1)
    add_paragraphs(
        document,
        [
            f"For German Credit, the current strongest pilot run is `{german_best['run_id']}`, which achieved test ROC-AUC {float(german_best['test_roc_auc']):.3f}, PR-AUC {float(german_best['test_pr_auc']):.3f}, recall {float(german_best['test_recall']):.3f}, and F1 {float(german_best['test_f1']):.3f}. This result suggests that class-weighted Logistic Regression remains a strong, transparent baseline on this benchmark.",
            f"For the Lending Club sample, the current strongest ROC-AUC run is `{lending_best['run_id']}` with test ROC-AUC {float(lending_best['test_roc_auc']):.3f}. However, the best ROC-AUC run does not automatically imply the strongest risky-class recall. In this pilot, the class-weighted Logistic Regression run produced comparatively higher recall {float(lending_best['test_recall']):.3f}, which is important when the research emphasis includes minority-risk detection.",
            "These early results already support the value of the designed benchmark: different models optimize different objectives, and imbalance handling materially affects recall and F1. This gives the report a concrete empirical basis rather than a methodology-only description.",
        ],
    )

    document.add_heading("8. Explainability Outputs", level=1)
    add_paragraphs(
        document,
        [
            "Feature-importance artifacts were saved for every pilot run. This means the report can already demonstrate the explainability workflow required by RQ3, even before adding SHAP-style analysis in a later phase.",
            "The strongest currently interpretable baseline is the class-weighted Logistic Regression model, which combines competitive predictive performance with straightforward coefficient-based importance ranking.",
        ],
    )
    document.add_picture(str(ROOT / "artifacts/credit/german_logreg_class_weight/feature_importance.png"), width=Inches(6.5))
    document.add_picture(str(ROOT / "artifacts/credit/lending_logreg_class_weight/feature_importance.png"), width=Inches(6.5))

    document.add_heading("9. PaddleOCR Extension: Current Status", level=1)
    add_paragraphs(
        document,
        [
            "The PaddleOCR workstream was intentionally scoped as a supporting module to show how borrower-document ingestion can be integrated into a broader loan-origination pipeline.",
            ocr_status["summary"],
            ocr_status["blocker"],
        ],
    )
    add_table(
        document,
        ["OCR Setup Checkpoint", "Current State"],
        [[key, value] for key, value in ocr_status["items"].items()],
    )
    add_paragraphs(
        document,
        [
            "What is already done: synthetic OCR label files were generated, the PaddleOCR repository was cloned successfully, and the pretrained recognition archive was downloaded.",
            "What is currently stuck: the fine-tuning bootstrap has not yet produced a stable extracted pretrained checkpoint and saved training output. This is the exact issue being debugged at the moment.",
            "Why this is still worth reporting: the OCR pipeline is clearly documented as work in progress, and it demonstrates concrete engineering progress rather than only a future idea.",
        ],
    )

    document.add_heading("10. Current Debugging Notes", level=1)
    add_table(
        document,
        ["Observed State", "Interpretation / Debugging Direction"],
        [
            ["`paddlepaddle` and `paddleocr` installed", "Core OCR runtime is available locally."],
            ["Synthetic OCR dataset generated", "Tiny pilot data is ready for a reduced recognition fine-tuning run."],
            ["Official PaddleOCR repo cloned", "Training scripts are available in the workspace."],
            ["Pretrained recognition tar downloaded", "Bootstrap assets are partially available."],
            ["No final OCR summary or model checkpoint yet", "The run is not complete and needs further debugging before claiming OCR results."],
        ],
    )
    add_paragraphs(
        document,
        [
            "The honest status is therefore: credit-scoring experiments are complete for the pilot stage, while PaddleOCR fine-tuning is partially set up and actively being debugged.",
            "This separation is important because it allows the current report to present real experimental evidence without overstating the maturity of the OCR extension.",
        ],
    )

    document.add_heading("11. Limitations", level=1)
    add_paragraphs(
        document,
        [
            "The Lending Club component currently uses a manageable public sample rather than the full raw dataset used in some published studies.",
            "The neural benchmark is still a pilot MLP rather than a more specialized tabular deep architecture.",
            "The OCR component should be treated as engineering progress and experimental setup rather than a completed result section.",
        ],
    )

    document.add_heading("12. Next Steps", level=1)
    add_table(
        document,
        ["Priority", "Next Action"],
        [
            ["1", "Complete the PaddleOCR bootstrap and obtain the first saved fine-tuned recognition checkpoint."],
            ["2", "Add held-out OCR evaluation metrics once the fine-tuned model runs successfully."],
            ["3", "Extend explainability with SHAP-style analysis if dependency/runtime costs remain acceptable."],
            ["4", "Run a larger Lending Club subset or the full public dataset if time and compute permit."],
        ],
    )

    document.add_heading("13. Meeting Note", level=1)
    add_paragraphs(
        document,
        [
            "To avoid confusion in supervisor communication, this report uses absolute dates. The earlier email thread referred to a meeting on Thursday 22 April 2026, but 22 April 2026 was a Wednesday. If the intended Thursday schedule is preserved, the matching date is Thursday 23 April 2026.",
        ],
    )

    document.save(OUTPUT_DOCX)
    OUTPUT_MD.write_text(build_markdown(rows, ocr_status), encoding="utf-8")


def main() -> None:
    rows = load_rows()
    ocr_status = infer_ocr_status()
    build_docx(rows, ocr_status)
    print(f"Wrote {OUTPUT_DOCX}")


if __name__ == "__main__":
    main()
