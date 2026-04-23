from __future__ import annotations

import csv
import json
from collections import defaultdict
from datetime import datetime
from pathlib import Path

from docx import Document


ROOT = Path("/Users/macbook/Downloads/Credit Scoring Model")
SUMMARY_CSV = ROOT / "artifacts/credit/benchmark_summary.csv"
OCR_SUMMARY = ROOT / "artifacts/ocr/paddle_ocr_pilot/ocr_run_summary.json"
REPORT_MD = ROOT / "report/SIT723_progress_report.md"
REPORT_DOCX = ROOT / "report/SIT723_progress_report.docx"
EMAIL_MD = ROOT / "report/supervisor_email_update.md"


def load_credit_summary() -> list[dict[str, str]]:
    if not SUMMARY_CSV.exists():
        raise FileNotFoundError(f"Missing benchmark summary: {SUMMARY_CSV}")
    with SUMMARY_CSV.open("r", encoding="utf-8") as handle:
        return list(csv.DictReader(handle))


def load_ocr_summary() -> dict:
    if not OCR_SUMMARY.exists():
        return {
            "status": "not_run",
            "message": "OCR experiment summary was not generated.",
        }
    return json.loads(OCR_SUMMARY.read_text(encoding="utf-8"))


def select_best_runs(rows: list[dict[str, str]]) -> dict[str, dict[str, str]]:
    grouped: dict[str, list[dict[str, str]]] = defaultdict(list)
    for row in rows:
        grouped[row["dataset"]].append(row)

    best = {}
    for dataset, dataset_rows in grouped.items():
        best[dataset] = max(dataset_rows, key=lambda row: float(row["test_roc_auc"]))
    return best


def build_markdown(rows: list[dict[str, str]], ocr_summary: dict) -> str:
    best_runs = select_best_runs(rows)
    today = datetime.now().strftime("%d %B %Y")
    lines = [
        "# SIT723 Progress Report",
        "",
        f"Prepared on {today}",
        "",
        "## Project focus",
        "",
        "This progress update keeps the approved research core on explainable AI and deep learning for credit scoring using German Credit and Lending Club data. A preliminary PaddleOCR document-intelligence module is included as a supporting loan-origination extension rather than a replacement for the core prediction study.",
        "",
        "## Research aim and questions",
        "",
        "The aim is to benchmark machine learning and deep learning approaches for credit-risk prediction while preserving robustness and practical explainability. The current pilot addresses three questions: how deep learning compares with strong tabular baselines, which imbalance strategies strengthen minority-class detection, and how explainability can be retained through feature-importance analysis.",
        "",
        "## Finalised methodology",
        "",
        "- Datasets: full German Credit data from UCI and a manageable public Lending Club sample sourced from a cleaned public CSV mirror.",
        "- Models: Logistic Regression, Random Forest, HistGradientBoosting, and MLP.",
        "- Data handling: train/validation/test split, median or mode imputation, one-hot encoding, scaling where required, and explicit imbalance strategies.",
        "- Robustness design: no balancing, class weighting, and random oversampling.",
        "- Metrics: ROC-AUC, PR-AUC, precision, recall, F1, accuracy, confusion matrix, and training time.",
        "- Explainability: feature-importance artifacts saved for every run.",
        "",
        "## Experimental design completed",
        "",
        "The pilot benchmark was implemented as a config-driven workflow so each run produces dataset metadata, metrics JSON, confusion matrix, ROC curve, PR curve, and feature-importance outputs. This ensures report claims can be traced back to saved artifacts rather than hand-entered values.",
        "",
        "## Preliminary experiments conducted",
        "",
    ]

    for dataset, best_row in best_runs.items():
        dataset_rows = [row for row in rows if row["dataset"] == dataset]
        lines.extend(
            [
                f"### {dataset}",
                "",
                f"- Pilot runs completed: {len(dataset_rows)}",
                f"- Best pilot ROC-AUC: {float(best_row['test_roc_auc']):.3f} using `{best_row['model']}` with `{best_row['imbalance_strategy']}`",
                f"- Best pilot PR-AUC: {float(best_row['test_pr_auc']):.3f}",
                f"- Best pilot recall on risky class: {float(best_row['test_recall']):.3f}",
                f"- Best pilot F1 on risky class: {float(best_row['test_f1']):.3f}",
                "",
            ]
        )

    lines.extend(
        [
            "## Early interpretation",
            "",
            "The pilot results are intended as preliminary evidence only. At this stage, the benchmark shows whether performance trends are consistent across a smaller public Lending Club sample and the classical German Credit benchmark, while also revealing how class-imbalance handling changes risky-borrower recall.",
            "",
            "## PaddleOCR module status",
            "",
        ]
    )

    lines.append(f"- OCR run status: `{ocr_summary.get('status', 'unknown')}`")
    lines.append(f"- OCR note: {ocr_summary.get('message', 'Synthetic recognition fine-tuning workflow executed.')}")
    lines.extend(
        [
            "",
            "## Limitations",
            "",
            "- The Lending Club component currently uses a public sample rather than the full paper-scale dataset.",
            "- The OCR component is framed as a preliminary module experiment with synthetic or small-scale data only.",
            "- Results should be treated as pilot findings rather than final thesis conclusions.",
            "",
            "## Next steps",
            "",
            "- Extend the Lending Club experiments to a larger public slice or the full dataset if compute and access permit.",
            "- Add calibration analysis and SHAP-style explanation if runtime dependencies are practical.",
            "- Improve the OCR fine-tuning pilot with more realistic scanned loan documents and field-level evaluation.",
            "- Consolidate the strongest pipeline into the final dissertation methodology chapter.",
            "",
            "## Meeting note",
            "",
            "The previous email thread mentioned a meeting at 12:00 PM on Thursday 22 April 2026, but 22 April 2026 was a Wednesday. This report therefore uses absolute dates and avoids repeating that inconsistency.",
        ]
    )
    return "\n".join(lines)


def build_email(rows: list[dict[str, str]], ocr_summary: dict) -> str:
    best_runs = select_best_runs(rows)
    german = best_runs.get("german_credit")
    lending = best_runs.get("lending_club_sample")
    lines = [
        "Subject: SIT723 progress update",
        "",
        "Dear Dr. Huda,",
        "",
        "Please find my SIT723 progress update below.",
        "",
        "I have now finalised the project methodology and experimental design for the current pilot phase. The study remains focused on benchmarking explainable AI and deep learning approaches for credit scoring using the German Credit dataset and a manageable public Lending Club sample.",
        "",
        "The current pilot experiments include Logistic Regression, Random Forest, Gradient Boosting, and MLP models, with comparisons across class weighting and oversampling strategies. Each run produces saved metrics and plots for traceability.",
        "",
        f"For the German Credit pilot, the strongest current run achieved ROC-AUC {float(german['test_roc_auc']):.3f}, PR-AUC {float(german['test_pr_auc']):.3f}, and risky-class recall {float(german['test_recall']):.3f}.",
        f"For the Lending Club pilot sample, the strongest current run achieved ROC-AUC {float(lending['test_roc_auc']):.3f}, PR-AUC {float(lending['test_pr_auc']):.3f}, and risky-class recall {float(lending['test_recall']):.3f}.",
        "",
        f"I also prepared the preliminary PaddleOCR workflow as a document-intelligence extension to the broader loan-origination context. The current OCR pilot status is: {ocr_summary.get('status', 'unknown')}.",
        "",
        "At this stage, these should be treated as preliminary pilot findings rather than final results. My next steps are to strengthen the comparison on a larger Lending Club slice, expand explainability analysis, and improve the OCR module with more realistic document samples.",
        "",
        "Please let me know if you would like me to prepare any specific additional material before our next meeting.",
        "",
        "Kind regards,",
        "",
        "Rudraksh Dhamija",
        "",
    ]
    return "\n".join(lines)


def write_docx(markdown_text: str) -> None:
    document = Document()
    for block in markdown_text.split("\n\n"):
        block = block.strip()
        if not block:
            continue
        if block.startswith("# "):
            document.add_heading(block[2:], level=1)
            continue
        if block.startswith("## "):
            document.add_heading(block[3:], level=2)
            continue
        if block.startswith("### "):
            document.add_heading(block[4:], level=3)
            continue
        if block.startswith("- "):
            for line in block.splitlines():
                document.add_paragraph(line[2:], style="List Bullet")
            continue
        document.add_paragraph(block)
    REPORT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    document.save(REPORT_DOCX)


def main() -> None:
    rows = load_credit_summary()
    ocr_summary = load_ocr_summary()
    markdown_text = build_markdown(rows, ocr_summary)
    REPORT_MD.parent.mkdir(parents=True, exist_ok=True)
    REPORT_MD.write_text(markdown_text, encoding="utf-8")
    EMAIL_MD.write_text(build_email(rows, ocr_summary), encoding="utf-8")
    write_docx(markdown_text)
    print(f"Wrote {REPORT_MD}")
    print(f"Wrote {REPORT_DOCX}")
    print(f"Wrote {EMAIL_MD}")


if __name__ == "__main__":
    main()
